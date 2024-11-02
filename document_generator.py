# document_generator.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        self.output_dir = "generated_documents"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def generate_document(self, structure):
        try:
            doc = Document()
            
            # Проверяем наличие необходимых полей в structure
            structure = {
                'type': structure.get('type', 'document'),
                'theme': structure.get('theme', 'Untitled'),
                'titlePage': structure.get('titlePage', False),
                'tableOfContents': structure.get('tableOfContents', False),
                'introduction': structure.get('introduction', False),
                'conclusion': structure.get('conclusion', False),
                'references': structure.get('references', False),
                'sections': structure.get('sections', []),
                'content': structure.get('content', '')
            }
            
            # Добавление титульной страницы
            if structure['titlePage']:
                self.add_title_page(doc, structure['theme'])
            
            # Добавление оглавления
            if structure['tableOfContents']:
                self.add_table_of_contents(doc)
            
            # Добавление введения
            if structure['introduction']:
                self.add_introduction(doc)
            
            # Добавление основных разделов
            for section in structure['sections']:
                self.add_section(doc, section)
            
            # Добавление основного контента
            doc.add_paragraph(structure['content'])
            
            # Добавление заключения
            if structure['conclusion']:
                self.add_conclusion(doc)
            
            # Добавление списка литературы
            if structure['references']:
                self.add_references(doc)
            
            # Создаем имя файла с временной меткой
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{structure['theme']}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Сохранение документа
            doc.save(filepath)
            
            return {
                "success": True,
                "message": f"Документ сохранен как {filename}"
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Ошибка при создании документа: {str(e)}"
            }

    def add_title_page(self, doc, theme):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run("НАЗВАНИЕ ОРГАНИЗАЦИИ\n")
        run.bold = True
        run.font.size = Pt(16)
        
        for _ in range(10):
            doc.add_paragraph()
        
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"Тема: {theme}\n")
        run.bold = True
        run.font.size = Pt(16)
        
        for _ in range(10):
            doc.add_paragraph()
        
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.add_run("Выполнил: ФИО\nГруппа: XXX\n")

    def add_table_of_contents(self, doc):
        doc.add_paragraph("Содержание")
        doc.add_page_break()

    def add_introduction(self, doc):
        doc.add_heading("Введение", level=1)
        doc.add_paragraph()

    def add_section(self, doc, section_name):
        doc.add_heading(section_name, level=1)
        doc.add_paragraph()

    def add_conclusion(self, doc):
        doc.add_heading("Заключение", level=1)
        doc.add_paragraph()

    def add_references(self, doc):
        doc.add_heading("Список литературы", level=1)
        doc.add_paragraph()